import os
from docx import Document
import json

# 验证Word文档质量
def validate_word_document(word_file):
    try:
        # 打开Word文档
        doc = Document(word_file)
        
        # 检查文档基本属性
        validation_result = {
            'file': os.path.basename(word_file),
            'status': 'success',
            'issues': [],
            'details': {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
        }
        
        # 检查段落内容
        empty_paragraphs = 0
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                empty_paragraphs += 1
        
        if empty_paragraphs > 10:
            validation_result['issues'].append(f'存在{empty_paragraphs}个空段落')
        
        # 检查表格
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            if rows < 2 or cols < 2:
                validation_result['issues'].append(f'表格{i+1}可能不完整')
        
        # 检查标题
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text)
        
        if len(headings) < 3:
            validation_result['issues'].append('文档标题结构可能不完整')
        
        # 检查文档长度
        if len(doc.paragraphs) < 50:
            validation_result['issues'].append('文档内容可能过于简短')
        
        return validation_result
        
    except Exception as e:
        return {
            'file': os.path.basename(word_file),
            'status': 'error',
            'issues': [f'无法打开文档: {str(e)}'],
            'details': {}
        }

# 批量验证
if __name__ == "__main__":
    # 目标目录
    target_dir = "d:\\王元元老师大创\\王元元老师（王诗涵）—软著申请单"
    
    # 要验证的文件
    files = [
        "软著申请表.docx",
        "用户使用说明书.docx",
        "代码统计报告.docx",
        "源代码清单.docx",
        "设计说明书.docx"
    ]
    
    # 验证结果
    validation_results = []
    
    # 验证每个文件
    for file in files:
        word_file = os.path.join(target_dir, file)
        
        if os.path.exists(word_file):
            result = validate_word_document(word_file)
            validation_results.append(result)
            print(f"验证完成: {file}")
        else:
            validation_results.append({
                'file': file,
                'status': 'error',
                'issues': ['文件不存在'],
                'details': {}
            })
            print(f"文件不存在: {file}")
    
    # 生成验证报告
    report = {
        'total_files': len(validation_results),
        'success_files': sum(1 for r in validation_results if r['status'] == 'success'),
        'error_files': sum(1 for r in validation_results if r['status'] == 'error'),
        'details': validation_results
    }
    
    # 保存验证报告
    report_file = os.path.join(target_dir, "质量验证报告.json")
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    # 打印验证结果摘要
    print("\n验证报告摘要:")
    print(f"总文件数: {report['total_files']}")
    print(f"成功: {report['success_files']}")
    print(f"错误: {report['error_files']}")
    print(f"验证报告已保存至: {report_file}")
