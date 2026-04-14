import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 优化Word文档样式
def optimize_word_document(word_file):
    # 打开Word文档
    doc = Document(word_file)
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    
    # 优化标题样式
    for paragraph in doc.paragraphs:
        # 检查是否是标题
        if paragraph.style.name.startswith('Heading'):
            # 设置标题字体和大小
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                if paragraph.style.name == 'Heading 1':
                    run.font.size = Pt(16)
                    run.font.bold = True
                elif paragraph.style.name == 'Heading 2':
                    run.font.size = Pt(14)
                    run.font.bold = True
                elif paragraph.style.name == 'Heading 3':
                    run.font.size = Pt(12)
                    run.font.bold = True
            # 标题居中对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 优化普通文本
        else:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(11)
            # 普通文本左对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 优化表格样式
    for table in doc.tables:
        # 表格居中对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 优化表头
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
        
        # 优化表格内容
        for row in table.rows[1:]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10)
    
    # 保存优化后的文档
    doc.save(word_file)
    print(f"优化完成: {word_file}")

# 批量优化
if __name__ == "__main__":
    # 目标目录
    target_dir = "d:\\王元元老师大创\\王元元老师（王诗涵）—软著申请单"
    
    # 要优化的文件
    files = [
        "软著申请表.docx",
        "用户使用说明书.docx",
        "代码统计报告.docx",
        "源代码清单.docx",
        "设计说明书.docx"
    ]
    
    # 优化每个文件
    for file in files:
        word_file = os.path.join(target_dir, file)
        
        if os.path.exists(word_file):
            optimize_word_document(word_file)
        else:
            print(f"文件不存在: {word_file}")
