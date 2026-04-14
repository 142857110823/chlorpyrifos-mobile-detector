import os
import markdown
from docx import Document
from docx.shared import Inches
import re

# 转换Markdown到Word
def markdown_to_word(markdown_file, word_file):
    # 读取Markdown文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 处理Markdown内容
    lines = markdown_text.split('\n')
    current_heading_level = 0
    
    for line in lines:
        line = line.strip()
        
        # 处理标题
        if line.startswith('#'):
            heading_level = line.count('#')
            heading_text = line.lstrip('#').strip()
            
            if heading_level == 1:
                doc.add_heading(heading_text, level=0)
            elif heading_level == 2:
                doc.add_heading(heading_text, level=1)
            elif heading_level == 3:
                doc.add_heading(heading_text, level=2)
            elif heading_level == 4:
                doc.add_heading(heading_text, level=3)
            elif heading_level == 5:
                doc.add_heading(heading_text, level=4)
            elif heading_level == 6:
                doc.add_heading(heading_text, level=5)
            
            current_heading_level = heading_level
        
        # 处理列表
        elif line.startswith('- '):
            # 查找或创建当前段落
            if not doc.paragraphs:
                doc.add_paragraph()
            
            # 添加项目符号
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line.lstrip('- '))
        
        # 处理表格
        elif line.startswith('|'):
            # 检查是否是表头
            if '|--------|------|' in line:
                continue
            
            # 解析表格行
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            
            # 如果是第一行，创建表格
            if not hasattr(doc, 'current_table') or doc.current_table is None:
                doc.current_table = doc.add_table(rows=1, cols=len(cells))
                hdr_cells = doc.current_table.rows[0].cells
                for i, cell in enumerate(cells):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = cell
            else:
                # 添加新行
                row_cells = doc.current_table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell
        
        # 处理代码块
        elif line.startswith('```'):
            # 跳过代码块标记
            continue
        
        # 处理普通文本
        elif line:
            # 检查是否是表格内容
            if hasattr(doc, 'current_table') and doc.current_table is not None:
                # 检查是否是表格结束
                if not line.startswith('|'):
                    doc.current_table = None
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)
        
        # 处理空行
        else:
            doc.add_paragraph()
    
    # 保存Word文档
    doc.save(word_file)
    print(f"转换完成: {markdown_file} -> {word_file}")

# 批量转换
if __name__ == "__main__":
    # 源目录
    source_dir = "d:\\王元元老师大创\\APP开发"
    # 目标目录
    target_dir = "d:\\王元元老师大创\\王元元老师（王诗涵）—软著申请单"
    
    # 确保目标目录存在
    os.makedirs(target_dir, exist_ok=True)
    
    # 要转换的文件
    files = [
        "软著申请表.md",
        "用户使用说明书.md",
        "代码统计报告.md",
        "源代码清单.md",
        "设计说明书.md"
    ]
    
    # 转换每个文件
    for file in files:
        markdown_file = os.path.join(source_dir, file)
        word_file = os.path.join(target_dir, file.replace('.md', '.docx'))
        
        if os.path.exists(markdown_file):
            markdown_to_word(markdown_file, word_file)
        else:
            print(f"文件不存在: {markdown_file}")
